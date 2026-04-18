"""
Word Generator - Xuất file DCCT theo định dạng Word chuẩn của Khoa CNTT - ĐH Đà Nẵng
Sử dụng python-docx để tạo tài liệu Word chuyên nghiệp
"""

import os
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path

from utils.logger import get_logger

logger = get_logger("export.word_generator")

OUTPUT_DIR = Path(__file__).parent.parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


async def export_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """
    Xuất DCCT ra file Word.
    
    Input state: final_dcct_data hoặc dữ liệu riêng lẻ
    Output state: export_ready=True, export_path
    """
    logger.info("[Export] Bắt đầu xuất file Word DCCT")

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError as e:
        logger.error(f"[Export] python-docx chưa được cài đặt: {e}")
        return {
            "export_ready": False,
            "errors": state.get("errors", []) + ["Export: python-docx chưa được cài đặt"],
        }

    # Lấy dữ liệu DCCT
    dcct_data = state.get("final_dcct_data") or _build_dcct_from_state(state)
    course_info = dcct_data.get("course_info", {})

    try:
        doc = Document()
        _set_document_style(doc)

        # ====== TRANG BÌA ======
        _add_cover_page(doc, course_info)

        # ====== 1. THÔNG TIN HỌC PHẦN ======
        _add_course_info_section(doc, course_info)

        # ====== 2. MÔ TẢ HỌC PHẦN ======
        _add_course_description(doc, course_info, state)

        # ====== 3. CHUẨN ĐẦU RA HỌC PHẦN (CLO) ======
        _add_clo_section(doc, dcct_data.get("clo_list", []))

        # ====== 4. MA TRẬN CLO - PLO ======
        _add_mapping_matrix_section(doc, dcct_data.get("clo_list", []), dcct_data.get("mapping_matrix", []))

        # ====== 5. KẾ HOẠCH GIẢNG DẠY ======
        _add_teaching_plan_section(doc, dcct_data.get("teaching_plan", []))

        # ====== 6. HỆ THỐNG ĐÁNH GIÁ ======
        _add_assessment_section(doc, dcct_data.get("assessment_plan", []))

        # ====== 7. RUBRIC ======
        _add_rubric_section(doc, dcct_data.get("rubrics", {}))

        # ====== LƯU FILE ======
        course_code = course_info.get("code", "DCCT")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"DCCT_{course_code}_{timestamp}.docx"
        filepath = OUTPUT_DIR / filename

        doc.save(str(filepath))
        logger.info(f"[Export] File xuất thành công: {filepath}")

        # ── Auto-index ĐCCT vào knowledge base Q&A ──────────────────────────
        try:
            from agents.qa_agent import index_dcct_from_state
            course_code_key = course_info.get("code", state.get("course_code", "UNKNOWN"))
            idx_info = index_dcct_from_state(course_code_key, state)
            logger.info(
                f"[Export] ĐCCT indexed vào Q&A store: "
                f"{idx_info['chunks_indexed']} chunks cho {course_code_key}"
            )
        except Exception as idx_err:
            logger.warning(f"[Export] Không thể index ĐCCT vào Q&A store: {idx_err}")

        return {
            "export_ready": True,
            "export_path": str(filepath),
        }

    except Exception as e:
        logger.error(f"[Export] Lỗi xuất Word: {e}")
        return {
            "export_ready": False,
            "errors": state.get("errors", []) + [f"Export: {e}"],
        }


def _set_document_style(doc):
    """Thiết lập style mặc định cho tài liệu."""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    import docx.oxml

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)


def _add_cover_page(doc, course_info: Dict):
    """Thêm trang bìa."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_paragraph()

    p = doc.add_paragraph("ĐẠI HỌC ĐÀ NẴNG")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 14, bold=True)

    p = doc.add_paragraph("TRƯỜNG ĐẠI HỌC BÁCH KHOA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 13, bold=True)

    p = doc.add_paragraph("KHOA CÔNG NGHỆ THÔNG TIN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 13, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph("ĐỀ CƯƠNG CHI TIẾT HỌC PHẦN")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 18, bold=True)

    doc.add_paragraph()

    course_name = course_info.get("name", "")
    course_code = course_info.get("code", "")
    p = doc.add_paragraph(f"{course_name}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 16, bold=True)

    p = doc.add_paragraph(f"({course_code})")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 14)

    doc.add_paragraph()
    doc.add_paragraph()

    year = datetime.now().year
    p = doc.add_paragraph(f"Đà Nẵng, năm {year}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(p, 12)

    doc.add_page_break()


def _add_course_info_section(doc, course_info: Dict):
    """Thêm phần thông tin học phần."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _add_heading(doc, "1. THÔNG TIN HỌC PHẦN", level=1)

    credits = course_info.get("credits", "3")
    theory = course_info.get("theory_periods", "N/A")
    lab = course_info.get("lab_periods", "N/A")

    info_items = [
        ("Tên học phần", course_info.get("name", "")),
        ("Mã học phần", course_info.get("code", "")),
        ("Số tín chỉ", f"{credits} tín chỉ ({theory} tiết LT + {lab} tiết TH)"),
        ("Loại học phần", course_info.get("course_type", "Bắt buộc")),
        ("Đối tượng", course_info.get("target_students", "Sinh viên CNTT")),
        ("Học phần tiên quyết", ", ".join(course_info.get("prerequisites", [])) or "Không có"),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(info_items):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        _set_cell_font(table.cell(i, 0), bold=True)

    doc.add_paragraph()


def _add_course_description(doc, course_info: Dict, state: Dict):
    """Thêm mô tả học phần."""
    _add_heading(doc, "2. MÔ TẢ HỌC PHẦN", level=1)

    summary = state.get("summary", course_info.get("description", "Không có mô tả."))
    p = doc.add_paragraph(summary)
    _set_para_font(p, 12)
    doc.add_paragraph()


def _add_clo_section(doc, clo_list: List[Dict]):
    """Thêm bảng CLO."""
    _add_heading(doc, "3. CHUẨN ĐẦU RA HỌC PHẦN (CLO)", level=1)

    if not clo_list:
        doc.add_paragraph("Chưa có CLO.")
        return

    headers = ["Mã CLO", "Mô tả", "Mức Bloom", "PI liên quan", "Mức IRMA"]
    table = doc.add_table(rows=1 + len(clo_list), cols=len(headers))
    table.style = "Table Grid"

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _set_cell_font(cell, bold=True, center=True)

    # Rows
    for row_idx, clo in enumerate(clo_list, 1):
        table.cell(row_idx, 0).text = clo.get("code", "")
        table.cell(row_idx, 1).text = clo.get("description", "")
        table.cell(row_idx, 2).text = clo.get("bloom_level_name", "N/A")
        table.cell(row_idx, 3).text = ", ".join(clo.get("pi_codes", []))
        table.cell(row_idx, 4).text = clo.get("mapping_level", "N/A")

    doc.add_paragraph()


def _add_mapping_matrix_section(doc, clo_list: List[Dict], mapping_matrix: List[Dict]):
    """Thêm ma trận CLO-PLO."""
    _add_heading(doc, "4. MA TRẬN ÁNH XẠ CLO - PLO", level=1)

    from utils.obe_utils import PLO_DATA

    if not clo_list:
        doc.add_paragraph("Chưa có dữ liệu mapping.")
        return

    # Lấy danh sách PLO được dùng
    used_plos = sorted(set(m.get("plo_code", "") for m in mapping_matrix if m.get("plo_code")))
    if not used_plos:
        doc.add_paragraph("Chưa có dữ liệu mapping.")
        return

    # Tạo lookup mapping
    mapping_lookup = {}
    for m in mapping_matrix:
        key = (m.get("clo_code", ""), m.get("plo_code", ""))
        mapping_lookup[key] = m.get("irma_level", "")

    # Bảng
    table = doc.add_table(rows=1 + len(clo_list), cols=1 + len(used_plos))
    table.style = "Table Grid"

    # Header row
    table.cell(0, 0).text = "CLO \\ PLO"
    _set_cell_font(table.cell(0, 0), bold=True, center=True)
    for j, plo in enumerate(used_plos, 1):
        table.cell(0, j).text = plo
        _set_cell_font(table.cell(0, j), bold=True, center=True)

    # Data rows
    for i, clo in enumerate(clo_list, 1):
        table.cell(i, 0).text = clo["code"]
        _set_cell_font(table.cell(i, 0), bold=True)
        for j, plo in enumerate(used_plos, 1):
            level = mapping_lookup.get((clo["code"], plo), "")
            table.cell(i, j).text = level

    doc.add_paragraph()

    # Chú thích
    p = doc.add_paragraph("Ghi chú: I=Introduce, R=Reinforce, M=Master, A=Apply")
    _set_para_font(p, 10, italic=True)
    doc.add_paragraph()


def _add_teaching_plan_section(doc, teaching_plan: List[Dict]):
    """Thêm kế hoạch giảng dạy."""
    _add_heading(doc, "5. KẾ HOẠCH GIẢNG DẠY", level=1)

    if not teaching_plan:
        doc.add_paragraph("Chưa có kế hoạch giảng dạy.")
        return

    headers = ["Buổi", "Tuần", "Loại", "Nội dung", "CLO", "Hoạt động", "Đánh giá"]
    table = doc.add_table(rows=1 + len(teaching_plan), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _set_cell_font(cell, bold=True, center=True)

    for row_idx, session in enumerate(teaching_plan, 1):
        table.cell(row_idx, 0).text = str(session.get("no", row_idx))
        table.cell(row_idx, 1).text = str(session.get("week", ""))
        table.cell(row_idx, 2).text = session.get("type", "LT")
        table.cell(row_idx, 3).text = session.get("content", "")
        table.cell(row_idx, 4).text = ", ".join(session.get("clo_codes", []))
        table.cell(row_idx, 5).text = session.get("activities", "")
        table.cell(row_idx, 6).text = session.get("assessment", "")

    doc.add_paragraph()


def _add_assessment_section(doc, assessment_plan: List[Dict]):
    """Thêm hệ thống đánh giá."""
    _add_heading(doc, "6. HỆ THỐNG ĐÁNH GIÁ", level=1)

    if not assessment_plan:
        doc.add_paragraph("Chưa có hệ thống đánh giá.")
        return

    headers = ["Cấu phần", "Tên", "Trọng số", "Hình thức", "CLO đánh giá"]
    table = doc.add_table(rows=1 + len(assessment_plan), cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _set_cell_font(cell, bold=True, center=True)

    total_weight = 0
    for row_idx, a in enumerate(assessment_plan, 1):
        weight = float(a.get("weight", 0))
        total_weight += weight
        table.cell(row_idx, 0).text = a.get("code", "")
        table.cell(row_idx, 1).text = a.get("name", "")
        table.cell(row_idx, 2).text = f"{weight * 100:.0f}%"
        table.cell(row_idx, 3).text = a.get("format", "")
        table.cell(row_idx, 4).text = ", ".join(a.get("clo_mapping", []))

    # Tổng trọng số
    doc.add_paragraph(f"Tổng trọng số: {total_weight * 100:.0f}%")
    doc.add_paragraph("Điểm đạt học phần: Điểm tổng ≥ 5.0/10")
    doc.add_paragraph()


def _add_rubric_section(doc, rubrics: Dict):
    """Thêm rubric đánh giá."""
    _add_heading(doc, "7. RUBRIC ĐÁNH GIÁ", level=1)

    if not rubrics:
        doc.add_paragraph("Chưa có rubric.")
        return

    for component_code, rubric in rubrics.items():
        _add_heading(doc, f"Rubric {component_code}", level=2)

        criteria = rubric.get("criteria", [])
        if not criteria:
            continue

        headers = ["Tiêu chí", "Trọng số", "Xuất sắc (90-100)", "Tốt (70-89)", "Đạt (50-69)", "Chưa đạt (<50)"]
        table = doc.add_table(rows=1 + len(criteria), cols=len(headers))
        table.style = "Table Grid"

        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            _set_cell_font(cell, bold=True, center=True)

        for row_idx, criterion in enumerate(criteria, 1):
            levels = criterion.get("levels", {})
            table.cell(row_idx, 0).text = criterion.get("criterion", "")
            table.cell(row_idx, 1).text = f"{criterion.get('weight_in_component', 0) * 100:.0f}%"
            table.cell(row_idx, 2).text = levels.get("excellent", {}).get("description", "")
            table.cell(row_idx, 3).text = levels.get("good", {}).get("description", "")
            table.cell(row_idx, 4).text = levels.get("pass", {}).get("description", "")
            table.cell(row_idx, 5).text = levels.get("fail", {}).get("description", "")

        doc.add_paragraph()


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def _add_heading(doc, text: str, level: int = 1):
    """Thêm heading với style phù hợp."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13 if level == 1 else 12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def _set_para_font(para, size: int = 12, bold: bool = False, italic: bool = False):
    """Thiết lập font cho paragraph."""
    from docx.shared import Pt
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    if not para.runs:
        run = para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def _set_cell_font(cell, bold: bool = False, center: bool = False):
    """Thiết lập font và alignment cho cell."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.bold = bold


def _build_dcct_from_state(state: Dict) -> Dict:
    """Xây dựng DCCT data từ state nếu final_dcct_data chưa có."""
    return {
        "course_info": {
            "code": state.get("course_code", ""),
            "name": state.get("course_name", ""),
            "credits": state.get("credits", "3"),
            **state.get("extracted_info", {}),
        },
        "clo_list": state.get("clo_list", []),
        "mapping_matrix": state.get("mapping_matrix", []),
        "teaching_plan": state.get("teaching_plan", []),
        "assessment_plan": state.get("assessment_plan", []),
        "rubrics": state.get("rubrics", {}),
        "confidence_score": state.get("confidence_score", 0),
    }
